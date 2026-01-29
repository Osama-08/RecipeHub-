"""
Convert markdown to Word document (.docx)
"""

import subprocess
import sys

# First, try to install required packages
packages = ['python-docx', 'markdown']

for package in packages:
    try:
        __import__(package.replace('-', '_'))
        print(f"✓ {package} already installed")
    except ImportError:
        print(f"Installing {package}...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", package, "-q"])
        print(f"✓ {package} installed")

# Now import and use
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

def markdown_to_docx(md_file, docx_file):
    """Convert markdown file to formatted Word document"""
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Process content line by line
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Headers
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            p = doc.add_heading(line[5:], level=4)
        
        # Horizontal rule
        elif line.startswith('---'):
            doc.add_paragraph('_' * 50)
        
        # Bullet lists
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:]
            # Check for checkboxes
            if text.startswith('[ ] '):
                text = '☐ ' + text[4:]
            elif text.startswith('[x] ') or text.startswith('[X] '):
                text = '☑ ' + text[4:]
            elif text.startswith('[/] '):
                text = '◐ ' + text[4:]
            
            p = doc.add_paragraph(text, style='List Bullet')
        
        # Numbered lists
        elif re.match(r'^\d+\. ', line):
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(text, style='List Number')
        
        # Code blocks
        elif line.startswith('```'):
            # Find end of code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            
            # Add code block
            code_text = '\n'.join(code_lines)
            p = doc.add_paragraph(code_text)
            p.style = 'No Spacing'
            for run in p.runs:
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
        
        # Tables
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # Parse table
            table_lines = [line]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].strip())
                i += 1
            i -= 1  # Back up one
            
            # Create table
            rows = [row.split('|')[1:-1] for row in table_lines if not set(row.replace('|', '').replace('-', '').strip()) == set()]
            if rows:
                table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                table.style = 'Light Grid Accent 1'
                
                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_data in enumerate(row_data):
                        table.rows[row_idx].cells[col_idx].text = cell_data.strip()
        
        # Block quotes
        elif line.startswith('> '):
            text = line[2:]
            # Handle multiple levels of quotes
            while text.startswith('> '):
                text = text[2:]
            
            p = doc.add_paragraph(text)
            p.paragraph_format.left_indent = Inches(0.5)
            for run in p.runs:
                run.italic = True
        
        # Regular paragraphs
        else:
            # Process inline formatting
            text = line
            
            # Remove markdown links but keep text - [text](url) -> text
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            
            p = doc.add_paragraph()
            
            # Split by formatting markers
            parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)', text)
            
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold
                    run = p.add_run(part[2:-2])
                    run.bold = True
                elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                    # Italic
                    run = p.add_run(part[1:-1])
                    run.italic = True
                elif part.startswith('`') and part.endswith('`'):
                    # Inline code
                    run = p.add_run(part[1:-1])
                    run.font.name = 'Consolas'
                    run.font.size = Pt(10)
                else:
                    # Regular text
                    p.add_run(part)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"\n✅ Successfully created: {docx_file}")

if __name__ == "__main__":
    md_file = "CLIENT_DOCUMENTATION.md"
    docx_file = "RecipeHub_Client_Documentation.docx"
    
    markdown_to_docx(md_file, docx_file)
    print(f"\n🎉 Word document ready!")
    print(f"📄 Location: {docx_file}")
